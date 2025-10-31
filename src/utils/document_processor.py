"""
Insyte AI - Document Processor
Extracts text from various document formats (PDF, DOCX, TXT).
"""

import logging
import os
from typing import Dict, Optional, Tuple
import io

class DocumentProcessor:
    """Process and extract text from different document formats."""
    
    def __init__(self):
        self.logger = logging.getLogger(__name__)
        self.supported_formats = ['.pdf', '.docx', '.txt', '.doc']
    
    def process_file(self, file_content: bytes, filename: str) -> Tuple[Optional[str], Dict]:
        """
        Process uploaded file and extract text content.
        
        Args:
            file_content: File content as bytes
            filename: Original filename
            
        Returns:
            Tuple of (extracted_text, metadata_dict)
        """
        file_ext = os.path.splitext(filename)[1].lower()
        metadata = {
            'filename': filename,
            'file_type': file_ext,
            'file_size': len(file_content)
        }
        
        try:
            if file_ext == '.pdf':
                text, extra_meta = self._process_pdf(file_content)
                metadata.update(extra_meta)
                return text, metadata
            elif file_ext in ['.docx', '.doc']:
                text, extra_meta = self._process_docx(file_content)
                metadata.update(extra_meta)
                return text, metadata
            elif file_ext == '.txt':
                text = self._process_txt(file_content)
                return text, metadata
            else:
                self.logger.warning(f"Unsupported file format: {file_ext}")
                return None, metadata
                
        except Exception as e:
            self.logger.error(f"Error processing {filename}: {str(e)}")
            metadata['error'] = str(e)
            return None, metadata
    
    def _process_pdf(self, file_content: bytes) -> Tuple[str, Dict]:
        """Extract text from PDF using multiple methods."""
        metadata = {}
        
        try:
            # Try pdfplumber first (better extraction)
            import pdfplumber
            
            with pdfplumber.open(io.BytesIO(file_content)) as pdf:
                text_parts = []
                metadata['page_count'] = len(pdf.pages)
                
                for page_num, page in enumerate(pdf.pages, 1):
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(f"--- Page {page_num} ---\n{page_text}")
                
                full_text = "\n\n".join(text_parts)
                metadata['extraction_method'] = 'pdfplumber'
                return full_text, metadata
                
        except ImportError:
            self.logger.info("pdfplumber not available, trying PyPDF2")
            
        try:
            # Fallback to PyPDF2
            import PyPDF2
            
            pdf_reader = PyPDF2.PdfReader(io.BytesIO(file_content))
            text_parts = []
            metadata['page_count'] = len(pdf_reader.pages)
            
            for page_num, page in enumerate(pdf_reader.pages, 1):
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(f"--- Page {page_num} ---\n{page_text}")
            
            full_text = "\n\n".join(text_parts)
            metadata['extraction_method'] = 'PyPDF2'
            return full_text, metadata
            
        except Exception as e:
            self.logger.error(f"PDF extraction failed: {str(e)}")
            raise
    
    def _process_docx(self, file_content: bytes) -> Tuple[str, Dict]:
        """Extract text from DOCX file."""
        try:
            from docx import Document
            
            doc = Document(io.BytesIO(file_content))
            
            # Extract paragraphs
            paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
            full_text = "\n\n".join(paragraphs)
            
            # Extract tables
            table_texts = []
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join([cell.text for cell in row.cells])
                    if row_text.strip():
                        table_texts.append(row_text)
            
            if table_texts:
                full_text += "\n\n--- Tables ---\n" + "\n".join(table_texts)
            
            metadata = {
                'paragraph_count': len(paragraphs),
                'table_count': len(doc.tables),
                'extraction_method': 'python-docx'
            }
            
            return full_text, metadata
            
        except Exception as e:
            self.logger.error(f"DOCX extraction failed: {str(e)}")
            raise
    
    def _process_txt(self, file_content: bytes) -> str:
        """Extract text from TXT file."""
        try:
            # Try UTF-8 first
            return file_content.decode('utf-8')
        except UnicodeDecodeError:
            # Fallback to latin-1
            try:
                return file_content.decode('latin-1')
            except Exception as e:
                self.logger.error(f"TXT decoding failed: {str(e)}")
                raise
    
    def is_supported(self, filename: str) -> bool:
        """Check if file format is supported."""
        file_ext = os.path.splitext(filename)[1].lower()
        return file_ext in self.supported_formats
    
    def get_file_info(self, file_content: bytes, filename: str) -> Dict:
        """Get basic file information without processing."""
        file_ext = os.path.splitext(filename)[1].lower()
        return {
            'filename': filename,
            'file_type': file_ext,
            'file_size': len(file_content),
            'file_size_mb': round(len(file_content) / (1024 * 1024), 2),
            'supported': self.is_supported(filename)
        }
